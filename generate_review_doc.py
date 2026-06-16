import re

from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt

from utils import strip_html


DEFAULT_FONT_SIZE = Pt(10)


def _resolve_highlight_color(color):
    if color is None:
        return None
    if isinstance(color, int):
        try:
            return WD_COLOR_INDEX(color)
        except Exception:
            return None
    if hasattr(color, "value"):
        return color
    if isinstance(color, str):
        upper = color.upper()
        if hasattr(WD_COLOR_INDEX, upper):
            return getattr(WD_COLOR_INDEX, upper)
    return None


def _apply_default_font(run):
    run.font.size = DEFAULT_FONT_SIZE


def _add_run_with_default_font(paragraph, text, bold=False):
    run = paragraph.add_run(str(text))
    run.bold = bold
    _apply_default_font(run)
    return run


def normalize_revamped_scenario(data):
    """
    Normalizes scenario payloads into a consistent dict shape.
    Accepts string or dict and returns:
    {
      "scenario_text": "..."
    }
    """
    if isinstance(data, dict):
        scenario_text = data.get("scenario") or data.get("scenario_text") or data.get("revamped_scenario") or ""
        return {"scenario_text": str(scenario_text)}

    return {"scenario_text": str(data or "")}


def _extract_plain_text(value):
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, dict):
        for key in ("scenario", "scenario_text", "revamped_scenario", "text", "value"):
            nested = value.get(key)
            if nested:
                return _extract_plain_text(nested)
        for nested in value.values():
            extracted = _extract_plain_text(nested)
            if extracted:
                return extracted
        return ""
    if isinstance(value, list):
        parts = [_extract_plain_text(item) for item in value]
        parts = [part for part in parts if part]
        return "\n".join(parts)
    return str(value)


def parse_blank_instructions(text):
    """
    Parses instruction blocks that may contain per-blank sections.
    Returns a list of printable instruction lines.
    """
    content = str(text or "").strip()
    if not content:
        return []

    pattern = re.compile(r"(Blank\s*\d+)\s*[:\-]\s*(.*?)(?=(?:\nBlank\s*\d+\s*[:\-])|$)", re.IGNORECASE | re.DOTALL)
    matches = list(pattern.finditer(content))
    if not matches:
        return [line for line in content.splitlines() if line.strip()]

    lines = []
    for match in matches:
        blank = match.group(1).title().replace("  ", " ")
        instruction = match.group(2).strip()
        lines.append(f"{blank}: {instruction}")
    return lines


def add_highlighted_paragraph(doc, text, highlight_color):
    paragraph = doc.add_paragraph()
    _add_run_with_default_font(paragraph, str(text))
    safe_color = _resolve_highlight_color(highlight_color)
    if safe_color is not None and paragraph.runs:
        paragraph.runs[0].font.highlight_color = safe_color


def add_plain_paragraphs(doc, text, highlight_color=None):
    content = _extract_plain_text(text)
    if not content.strip():
        paragraph = doc.add_paragraph("")
        for run in paragraph.runs:
            _apply_default_font(run)
        return

    for line in content.splitlines():
        paragraph = doc.add_paragraph()
        _add_run_with_default_font(paragraph, line)
        safe_color = _resolve_highlight_color(highlight_color)
        if safe_color is not None and paragraph.runs:
            paragraph.runs[0].font.highlight_color = safe_color


def _get_section_text(mapped, qc_result, approved, section_name):
    if section_name == "scenario":
        if approved.get(section_name):
            revamped = qc_result.get("scenario_qc", {}).get("revamped_scenario")
            normalized = normalize_revamped_scenario(revamped if isinstance(revamped, dict) else {"scenario": revamped})
            text = normalized.get("scenario_text") or qc_result.get("revamped_scenario") or mapped.get("scenario", "")
        else:
            text = mapped.get("scenario", "")
    elif section_name == "instructions":
        text = qc_result.get("revamped_instructions", "") if approved.get(section_name) else mapped.get("instructions", "")
    elif section_name == "code_snippet":
        text = qc_result.get("revamped_code_snippet", "") if approved.get(section_name) else mapped.get("code_snippet", "")
    else:
        text = mapped.get(section_name, "")

    text = _extract_plain_text(text)
    if "<" in text:
        text = strip_html(text)
    return text


def _strip_instructions_from_scenario(scenario_text):
    content = str(scenario_text or "")
    marker = re.search(r"(?i)(Complete the code as per the given instructions:)", content)
    if marker:
        content = content[:marker.start()].strip()
    return content


def _add_instruction_paragraphs(doc, original_lines, revamped_lines, approved):
    added = set()
    for idx, orig_line in enumerate(original_lines):
        if idx < len(revamped_lines) and str(revamped_lines[idx] or "").strip() != str(orig_line or "").strip():
            add_plain_paragraphs(doc, revamped_lines[idx], highlight_color=WD_COLOR_INDEX.YELLOW)
            added.add(idx)

    for idx, line in enumerate(revamped_lines):
        if idx not in added:
            add_plain_paragraphs(doc, line, highlight_color=None)


def add_question_section(doc, que_id, mapped, qc_result, approved):
    topic = mapped.get("topic", "")
    subtopic = mapped.get("sub_topic", "")
    difficulty = mapped.get("difficulty", "")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(f"Q Id: {que_id}")
    run.bold = True
    run.font.size = DEFAULT_FONT_SIZE

    add_highlighted_paragraph(doc, f"Topic: {topic}", highlight_color=WD_COLOR_INDEX.GRAY_25)
    add_highlighted_paragraph(doc, f"Sub-topic: {subtopic}", highlight_color=WD_COLOR_INDEX.GRAY_25)
    add_highlighted_paragraph(doc, f"Difficulty: {difficulty}", highlight_color=WD_COLOR_INDEX.GRAY_25)
    doc.add_paragraph("")

    scenario_changed = bool(approved.get("scenario"))
    instructions_changed = bool(approved.get("instructions"))
    code_changed = bool(approved.get("code_snippet"))
    alt_changed = bool(approved.get("alternate_answers"))

    if scenario_changed:
        add_highlighted_paragraph(doc, "Change in Scenario", highlight_color=WD_COLOR_INDEX.YELLOW)
    if code_changed:
        add_highlighted_paragraph(doc, "Change in Code Snippet", highlight_color=WD_COLOR_INDEX.YELLOW)
    if instructions_changed:
        add_highlighted_paragraph(doc, "Change in Instructions", highlight_color=WD_COLOR_INDEX.YELLOW)
    if alt_changed:
        add_highlighted_paragraph(doc, "Change in alternative answers", highlight_color=WD_COLOR_INDEX.YELLOW)

    if not any([scenario_changed, code_changed, instructions_changed, alt_changed]):
        add_highlighted_paragraph(doc, "No changes applied - original content retained", highlight_color=WD_COLOR_INDEX.BRIGHT_GREEN)

    doc.add_paragraph("")

    scenario_text = _get_section_text(mapped, qc_result, approved, "scenario")
    scenario_text = _strip_instructions_from_scenario(scenario_text)
    instructions_text = _get_section_text(mapped, qc_result, approved, "instructions")
    code_text = _get_section_text(mapped, qc_result, approved, "code_snippet")

    p = doc.add_paragraph()
    run = p.add_run("Question:")
    _apply_default_font(run)
    add_plain_paragraphs(doc, scenario_text, highlight_color=WD_COLOR_INDEX.YELLOW if scenario_changed else None)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Complete the code as per the given instructions:")
    _apply_default_font(run)

    original_instructions = mapped.get("instructions", "")
    original_lines = parse_blank_instructions(original_instructions) or ([original_instructions] if original_instructions else [""])
    revamped_lines = parse_blank_instructions(instructions_text) or ([instructions_text] if instructions_text else [""])

    if instructions_changed:
        _add_instruction_paragraphs(doc, original_lines, revamped_lines, approved)
    else:
        add_plain_paragraphs(doc, "\n".join(revamped_lines), highlight_color=None)
    doc.add_paragraph("")

    doc.add_paragraph(
        "A few lines in the Sample Script are missing (Enter your code here). "
        "You need to complete the code as per the given instructions."
    )
    doc.add_paragraph("")

    doc.add_paragraph("Sample Script:")
    doc.add_paragraph("")
    add_plain_paragraphs(
        doc,
        strip_html(code_text) if "<" in str(code_text) else code_text,
        highlight_color=WD_COLOR_INDEX.YELLOW if code_changed else None,
    )
    doc.add_paragraph("")

    doc.add_paragraph("Answers:")
    answers = mapped.get("answers", {})
    for blank_key in sorted(answers.keys()):
        doc.add_paragraph(f"{blank_key}: {answers[blank_key]}")
    doc.add_paragraph("")

    if alt_changed:
        alt_answers = qc_result.get("alternate_answers", {})
        doc.add_paragraph("--- Alternate Answers ---")
        doc.add_paragraph("")
        for blank_key in sorted(alt_answers.keys()):
            doc.add_paragraph(f"{blank_key}:")
            for i, value in enumerate(alt_answers.get(blank_key, []), 1):
                doc.add_paragraph(f"{i}. {value}")
            doc.add_paragraph("")


def _ensure_default_font(doc):
    try:
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = DEFAULT_FONT_SIZE
    except Exception:
        pass


def build_review_document_single(que_id, mapped, qc_result, approved, output_path):
    from docx import Document

    doc = Document()
    _ensure_default_font(doc)
    doc.add_paragraph("AILB QC & Revamp Review")
    doc.add_paragraph("1 Question")
    doc.add_paragraph("1")

    add_question_section(doc, que_id, mapped, qc_result, approved)

    doc.save(output_path)
    return output_path


def build_review_document_multi(questions_data, output_path):
    """
    Build a single DOCX containing all reviewed questions.
    questions_data should be a list of dicts with keys:
    {que_id, mapped, qc_result, approved}
    """
    from docx import Document

    doc = Document()
    _ensure_default_font(doc)
    doc.add_paragraph("AILB QC & Revamp Review")
    doc.add_paragraph(f"{len(questions_data)} Questions")
    doc.add_paragraph(str(len(questions_data)))

    for index, question_data in enumerate(questions_data):
        if index > 0:
            doc.add_page_break()
        add_question_section(
            doc,
            question_data.get("que_id", ""),
            question_data.get("mapped", {}),
            question_data.get("qc_result", {}),
            question_data.get("approved", {}),
        )

    doc.save(output_path)
    return output_path
